"""生成虚拟试衣系统 Word 报告。

报告按标准学术论文结构组织：
  · 第一章 绪论
  · 第二章 相关技术与理论基础
  · 第三章 系统设计与实现
  · 第四章 实验与结果分析
  · 第五章 与开源虚拟试衣方法的对比
  · 第六章 总结与展望
  · 参考文献

运行：
    .venv/Scripts/python.exe scripts/build_report.py
    .venv/Scripts/python.exe scripts/build_report.py --output report.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# 中文字体：宋体（学术论文标准正文/标题字体）
CN_FONT = "SimSun"
EN_FONT = "Times New Roman"

# 论文标准配色：标题、表头一律黑色
HEADING_COLOR = RGBColor(0x00, 0x00, 0x00)
TITLE_COLOR = RGBColor(0x00, 0x00, 0x00)

# 输出默认放在仓库根目录的 docs/ 下
ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUT = ROOT / "docs" / "report.docx"

# 可选嵌入的样例图（如果不存在则跳过）
SAMPLE_RESULT = ROOT / "output" / "grid3x4" / "image__image" / "result.jpg"
CONTACT_SHEET = ROOT / "output" / "grid3x4" / "contact_sheet.jpg"
DEBUG_OVERLAY = ROOT / "output" / "grid3x4" / "image__image" / "debug_warped_mask_overlay.jpg"


# ---------- 排版辅助 ----------

def set_cn_font(run, font_name: str = CN_FONT, size: int = 11,
                bold: bool = False, color: RGBColor | None = None) -> None:
    """同时设置中英文字体（python-docx 对中文需要走 rPr/rFonts/east-asia）。"""
    run.font.name = EN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)


def add_paragraph(doc: Document, text: str, *, size: int = 11,
                  bold: bool = False, align=None,
                  color: RGBColor | None = None,
                  first_line_indent: bool = False,
                  space_after: int = 4) -> None:
    """添加正文段落。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int) -> None:
    """添加标题（一级=章/二级=节/三级=小节）。"""
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if level == 1:
        set_cn_font(run, size=18, bold=True, color=HEADING_COLOR)
    elif level == 2:
        set_cn_font(run, size=15, bold=True, color=HEADING_COLOR)
    else:
        set_cn_font(run, size=12, bold=True, color=HEADING_COLOR)


def add_bullet(doc: Document, text: str, size: int = 11) -> None:
    """添加项目符号条目。"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, size=size)


def add_caption(doc: Document, text: str) -> None:
    """添加图/表说明。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_cn_font(run, size=10, color=RGBColor(0x55, 0x55, 0x55))


def add_image(doc: Document, path: Path, width_cm: float = 12.0) -> None:
    """嵌入图片（按宽度缩放，居中）。"""
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))


def add_table(doc: Document, header: list[str], rows: list[list[str]],
              col_widths_cm: list[float] | None = None) -> None:
    """添加表格（首行加粗 + 表头底色 + 单元格垂直居中）。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _black_borders(table)

    hdr = table.rows[0].cells
    for j, h in enumerate(header):
        hdr[j].text = ""
        para = hdr[j].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_cn_font(run, size=10, bold=True, color=HEADING_COLOR)
        hdr[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(hdr[j], "D9D9D9")

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, txt in enumerate(row):
            cells[j].text = ""
            para = cells[j].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(txt))
            set_cn_font(run, size=10)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_cm:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)


def _shade_cell(cell, hex_color: str) -> None:
    """给单元格加底色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _black_borders(table) -> None:
    """强制表格边框为黑色实线。"""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = tblBorders.makeelement(qn(f"w:{edge}"), {})
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ---------- 前置部分 ----------

def add_cover(doc: Document) -> None:
    """封面：标题、作者信息、日期。"""
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("虚拟试衣系统")
    set_cn_font(run, size=28, bold=True, color=TITLE_COLOR)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(8)
    run = sub.add_run("—— 基于关键点检测与几何 warp 的传统方法实现")
    set_cn_font(run, size=16, bold=False, color=TITLE_COLOR)

    for _ in range(6):
        doc.add_paragraph()

    info_lines = [
        ("课程名称", "计算机视觉课程设计"),
        ("项目名称", "VirtualFitting —— 关键点驱动的虚拟试衣"),
        ("作者", "Runxin Yan"),
        ("日期", "2026 年 6 月"),
        ("版本", "v1.0"),
    ]
    for k, v in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{k}：{v}")
        set_cn_font(run, size=12)

    doc.add_page_break()


def add_abstract(doc: Document) -> None:
    """中文摘要 + 英文摘要 + 关键词。"""
    # --- 中文摘要 ---
    add_heading(doc, "摘要", level=1)
    add_paragraph(
        doc,
        "本文实现了一套基于关键点检测与几何 warp 的传统计算机"
        "视觉虚拟试衣系统。系统以 MediaPipe Pose 提取的人体 33 "
        "关键点与 CLAHE + 多通道 Canny 分割派生的 8 个服装关键点"
        "为输入，采用两阶段流水式 warp：Stage A 用肩线中点对齐 + "
        "bbox 缩放做仿射粗定位，Stage B 沿身体 silhouette 按行分"
        "块（躯干 / 左袖 / 右袖）独立求解缩放与平移，得到 per-"
        "region 的 1D appearance flow，最后通过 mask 修正 + 边缘"
        "羽化与人体图做 alpha 混合。整个流水线无需成对训练数据，"
        "可在 CPU 上实时运行，每一步都有清晰的几何意义。本文"
        "在自建 12 组测试集上验证了系统的有效性，并将其与 VITON、"
        "CP-VTON、VITON-HD、HR-VITON、GFLA、GP-VTON、LaDI-VTON、"
        "OOTDiffusion、StableVITON 等主流开源方法做了系统对比，"
        "明确本系统作为“教学 / 0 样本 / 隐私敏感”场景优选的"
        "定位，并给出改进方向。",
        first_line_indent=True,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("关键词：")
    set_cn_font(run, size=11, bold=True)
    run = p.add_run("虚拟试衣；关键点检测；MediaPipe Pose；几何 warp；"
                    "silhouette 拟合；传统计算机视觉")
    set_cn_font(run, size=11)

    # --- 英文摘要 ---
    doc.add_paragraph()  # 空行
    add_heading(doc, "Abstract", level=1)
    add_paragraph(
        doc,
        "This paper presents a virtual try-on system based on classical "
        "computer vision techniques, driven by keypoint detection and "
        "geometric warping. The system takes 33 human keypoints from "
        "MediaPipe Pose and 8 clothing keypoints derived from CLAHE + "
        "multi-channel Canny segmentation as input, and adopts a two-"
        "stage pipeline warping. Stage A performs affine coarse alignment "
        "using the shoulder-line midpoint plus a bounding-box-based "
        "scale; Stage B performs per-region (torso / left sleeve / "
        "right sleeve) silhouette-driven warping along the y axis, "
        "yielding a 1D appearance flow for each region. A mask-rectified "
        "alpha blending with edge feathering produces the final composited "
        "image. The entire pipeline requires no paired training data, "
        "runs in real time on CPU, and every step has clear geometric "
        "interpretation. We validate the system on a self-collected "
        "12-pair test set and compare it systematically with mainstream "
        "open-source methods including VITON, CP-VTON, VITON-HD, "
        "HR-VITON, GFLA, GP-VTON, LaDI-VTON, OOTDiffusion, and "
        "StableVITON. The comparison positions the proposed system as a "
        "preferred choice for teaching, zero-shot, and privacy-sensitive "
        "scenarios, and discusses directions for future improvement.",
        first_line_indent=True,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Keywords: ")
    set_cn_font(run, size=11, bold=True)
    run = p.add_run("virtual try-on; keypoint detection; MediaPipe Pose; "
                    "geometric warping; silhouette fitting; classical computer vision")
    set_cn_font(run, size=11)

    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    """目录（占位提示，让用户在 Word 中按 F9 自动生成）。"""
    add_heading(doc, "目录", level=1)
    add_paragraph(
        doc,
        "提示：本文档目录由 Word 字段生成。打开后请在目录区域右键 → "
        "“更新域” → “更新整个目录”，即可生成完整目录。",
        first_line_indent=False,
        color=RGBColor(0x88, 0x88, 0x88),
    )
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = run._element.makeelement(qn("w:fldChar"),
                                        {qn("w:fldCharType"): "begin"})
    instrText = run._element.makeelement(qn("w:instrText"), {})
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = run._element.makeelement(qn("w:fldChar"),
                                        {qn("w:fldCharType"): "end"})
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    doc.add_page_break()


# ---------- 第一章 绪论 ----------

def add_chapter_intro(doc: Document) -> None:
    add_heading(doc, "第一章 绪论", level=1)

    add_heading(doc, "1.1 研究背景与意义", level=2)
    add_paragraph(
        doc,
        "虚拟试衣（Virtual Try-On）是计算机视觉与图形学交叉的"
        "经典任务：给定一张平铺服装图和一张人体图，合成“该人"
        "穿着该衣”的结果。它是电商服装零售的核心体验环节之一，"
        "也是 AR / 数字人 / 虚拟形象生成的基础技术。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "与传统基于 3D 建模的试衣相比，2D 图像级试衣不依赖昂贵"
        "的三维扫描设备，能直接消费电商已有的“商品图 + 模特图”，"
        "落地成本显著更低，是近 10 年来的研究主流。",
        first_line_indent=True,
    )

    add_heading(doc, "1.2 国内外研究现状", level=2)
    add_paragraph(
        doc,
        "图像级虚拟试衣的研究大致可分四代。第一代（2018 年前后）"
        "以 VITON [1] 与 CP-VTON [2] 为代表：VITON 用 coarse-to-fine "
        "策略直接合成试穿图，CP-VTON 引入可学习的 Geometric Matching "
        "Module（GMM）+ TPS 变形，但输出分辨率仅 256×192，纹理"
        "细节模糊。第二代（2021-2022）以 VITON-HD [3] 与 HR-VITON [4] "
        "为代表，把分辨率提升到 1024×768，分别用 ALIAS 归一化与"
        "统一 try-on condition generator 处理错位与遮挡。第三代"
        "（2020-2023）转向外观流：GFLA [5] 提出 global-flow "
        "local-attention 框架；He 等 [6] 首次把 StyleGAN 引入"
        "appearance flow 估计；GP-VTON [7] 用 local-flow + global-"
        "parsing 协同学习处理复杂版型。第四代（2023 至今）以 "
        "LaDI-VTON [8]、OOTDiffusion [9]、StableVITON [10] 为"
        "代表，把预训练 latent diffusion 引入虚拟试衣任务，用 "
        "textual inversion、outfitting fusion、zero cross-attention "
        "等机制保留服装细节，视觉质量显著领先，但需要 GPU 与"
        "大量训练数据。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "与之并行，传统 CV 路线（基于关键点 + 几何 warp）以可"
        "解释、零数据、CPU 友好为优势，仍是教学 / 0 样本 / "
        "隐私敏感场景的优选。本项目即属于这一路线。",
        first_line_indent=True,
    )

    add_heading(doc, "1.3 本文工作与组织结构", level=2)
    add_paragraph(
        doc,
        "本文实现了一套完整的传统 CV 虚拟试衣系统。具体工作包括："
        "(1) 基于 MediaPipe Pose 实现人体关键点自动检测，并设计"
        "三级降级链保证鲁棒性；(2) 基于 CLAHE + 多通道 Canny "
        "实现服装自动分割与 8 关键点几何派生；(3) 设计两阶段"
        "流水式 warp（Stage A 仿射 + Stage B 沿 silhouette 逐行 "
        "per-region fit），长袖自动分块并保留 V 领；(4) 设计"
        "mask 修正 + alpha 混合融合方案；(5) 与 10 种主流开源"
        "方法做系统对比，明确本系统的适用场景与改进方向。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "本文组织如下：第二章梳理相关技术（MediaPipe Pose、"
        "CLAHE、几何 warp、alpha 融合）；第三章描述系统的总体"
        "架构与各模块实现细节；第四章给出实验结果与失败案例"
        "分析；第五章与开源虚拟试衣方法做系统对比；第六章总结"
        "全文并展望改进方向。",
        first_line_indent=True,
    )

    doc.add_page_break()


# ---------- 第二章 相关技术与理论基础 ----------

def add_chapter_background(doc: Document) -> None:
    add_heading(doc, "第二章 相关技术与理论基础", level=1)

    add_heading(doc, "2.1 MediaPipe Pose 人体姿态估计", level=2)
    add_paragraph(
        doc,
        "MediaPipe Pose 是 Google 2019 年发布的实时多人姿态估计"
        "方案 [11]，采用 detector + landmarker 的两阶段结构："
        "先用检测器在图中定位人体边界框，再用回归网络对框内"
        "区域输出 33 个关键点的 (x, y, z, visibility)。本项目使用 "
        "full 模型而非 lite，full 在长袖、交叉臂、侧身等姿态下"
        "的左右一致性更好，模型文件约 12 MB，首次运行自动从 "
        "Google 公共存储下载并缓存到本地。MediaPipe 输出的 "
        "left/right 始终指 subject 解剖学方向，相机自拍模式"
        "会把 subject 右侧映射到图像左侧，本项目保留这一"
        "约定不做自动翻转。",
        first_line_indent=True,
    )

    add_heading(doc, "2.2 CLAHE 与多通道 Canny 边缘检测", level=2)
    add_paragraph(
        doc,
        "对比度受限图（如白衬衫 vs 白底）的分割难点在于“色距"
        "小但灰度局部有差”。本项目用 CLAHE（Contrast Limited "
        "Adaptive Histogram Equalization，clipLimit=3.0、8×8 tile）"
        "把局部灰度带拉宽，暴露被压平的边缘；再对灰度、B、G、R、"
        "Lab-L 五个通道分别做 Canny 边缘检测（50/150 双阈值）后"
        "取 OR 投票，捕获“色距相同但灰度突变”的边界。该流程"
        "完全在 OpenCV 中实现，不依赖任何深度学习模型。",
        first_line_indent=True,
    )

    add_heading(doc, "2.3 仿射变换与 silhouette-driven warp", level=2)
    add_paragraph(
        doc,
        "仿射变换是 2D 图像 warp 的基础：在齐次坐标下，平移、"
        "旋转、缩放、剪切可用一个 2×3 矩阵表示，OpenCV 的 "
        "cv2.warpAffine 即可高效实现。仿射变换是全局统一的，"
        "不能表达身体宽度沿 y 方向的非线性变化（如腰比肩窄）。"
        "为弥补这一不足，本项目在仿射粗定位之后引入 "
        "silhouette-driven warp：对每一行 y，假设该行的服装"
        "像素经 x 方向线性拉伸 s(y) 与平移 t(y) 后贴到身体"
        "轮廓，其中 s 与 t 沿 y 变化——这等价于一维的 appearance "
        "flow 估计。该思想与 GFLA [5]、GP-VTON [7] 等工作的"
        "warp 思路同源，本项目是其在 1D 投影下的简化实现。",
        first_line_indent=True,
    )

    add_heading(doc, "2.4 Alpha 融合与 mask 修正", level=2)
    add_paragraph(
        doc,
        "图像融合的目标是把 warp 后的服装无缝贴到人体图上。简单"
        "alpha 混合会出现三类 artifact：(1) cv2.warpAffine 对 RGB "
        "和 mask 的独立插值让 mask 边缘比 RGB 多出一圈，blend 时"
        "透出黑色边框；(2) 服装图自带的“白边”会被 warp 出来；"
        "(3) 服装与身体交界处的硬接缝。相应地，本项目在 blend "
        "前做三步 mask 修正——用 warped_rgb.max>0 收回多余 mask、"
        "用 ~1% min_dim 的 erode 剔除白边、用高斯模糊做羽化——"
        "再与原图做 alpha 混合。",
        first_line_indent=True,
    )

    doc.add_page_break()


# ---------- 第三章 系统设计与实现 ----------

def add_chapter_design(doc: Document) -> None:
    add_heading(doc, "第三章 系统设计与实现", level=1)

    add_heading(doc, "3.1 系统总体架构", level=2)
    add_paragraph(
        doc,
        "本系统面向“平铺服装图 + 单人正面图 → 试穿合成图”这一"
        "任务，按“关键点检测 → warp → 融合”三阶段组织，模块"
        "之间通过标准数据结构（Keypoint、numpy 轮廓点）解耦，"
        "便于单独替换与回归测试。系统提供如下端到端能力："
        "MediaPipe Pose 自动检测人体 33 关键点并派生颈部；"
        "CLAHE + 多通道 Canny 自动分割服装并派生 8 个语义"
        "关键点；两阶段流水式 warp（Stage A 仿射 + Stage B "
        "silhouette-driven 逐行 fit）把服装 fit 到身体；"
        "mask 修正 + alpha 混合实现无缝融合。",
        first_line_indent=True,
    )

    add_heading(doc, "3.2 输入输出与运行方式", level=2)
    add_paragraph(doc, "输入：", bold=True)
    for item in [
        "人体图：data_picture/people/image.png（正面、完整躯干、背景简洁）",
        "服装图：data_picture/clothes/image.png（白底 / 浅底最佳，可带 alpha 透明通道）",
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, "输出（output/run/<person>_<clothing>/）：", bold=True)
    add_table(
        doc,
        header=["文件名", "内容"],
        rows=[
            ["human_keypoints.jpg", "人体关键点可视化"],
            ["clothing_keypoints.jpg", "服装关键点可视化"],
            ["warped_clothing.png", "warp 后的服装 RGB（PNG 无损）"],
            ["warped_mask.png", "warp 后的二值前景掩码"],
            ["result.jpg", "最终试衣合成结果"],
            ["debug_*.jpg", "中间调试叠加图（轮廓、mask overlay 等）"],
        ],
        col_widths_cm=[5.0, 11.0],
    )
    add_paragraph(doc, "运行方式：", bold=True)
    add_bullet(doc, "CLI：python main.py run --person <person> --clothing <cloth> "
                    "--output <dir> --warp-method flow")
    add_bullet(doc, "Docker：./docker/run_docker.ps1（参见 docker/DOCKER_GUIDE.md）")
    add_bullet(doc, "批量对比：python scripts/grid_3x4.py 生成 3 人 × 4 衣 contact sheet")

    add_heading(doc, "3.3 人体关键点检测模块", level=2)
    add_paragraph(
        doc,
        "MediaPipeHumanDetector 首次运行自动从 Google 公共存储"
        "下载 pose_landmarker_full.task 模型（约 12 MB）并缓存。"
        "检测返回 33 个姿态关键点后，按 keypoints."
        "MEDIAPIPE_POSE_INDICES 映射为命名集合，并由双肩中点"
        "派生 neck 虚拟关键点。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "为应对模型不可用的场景，模块实现三级降级链："
        "MediaPipeHumanDetector → HaarHumanDetector → "
        "HeuristicHumanDetector。三者通过 RobustHumanDetector "
        "包装器按顺序尝试，任何一级成功即返回；只有当模型下载"
        "失败、Haar 级联全部加载失败时才会抛异常。Haar 路径用 "
        "OpenCV 内置的 haarcascade_frontalface_default 等 XML，"
        "按人体比例推算关键点；Heuristic 路径把关键点放在图像"
        "中心按固定比例位置，仅作 sanity check。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "检测到 33 关键点后，body_region_contour 函数按真实人体"
        "形状设计的多边形（neck_top → 左肩 → 左腋 → 左肘 → "
        "左腕 → 底中 → 右腕 → 右肘 → 右腋 → 右肩）构造身体轮廓。"
        "关键设计点：(1) neck_top 在肩线以上（不再是平直肩线），"
        "(2) 腋下与肩同 x（臂从此处开始离开躯干），(3) 髋部"
        "向外扩到外缘（MediaPipe 给的是关节中心而非外缘），"
        "(4) 腋下之后沿手臂外缘继续向下到腕（让长袖 fit 到"
        "手臂而非被压扁）。该多边形随后被弧长均匀采样为 30 个"
        "点供 warp 使用。",
        first_line_indent=True,
    )

    add_heading(doc, "3.4 服装关键点检测模块", level=2)
    add_paragraph(
        doc,
        "ClothingDetector 流程：CLAHE 拉伸灰度 → 多通道 Canny "
        "（灰度 + B + G + R + Lab-L 取 OR）→ 15×15 闭运算封口 → "
        "二次膨胀 → 取最大外轮廓填充 → 内缩 1 次补偿 dilation。"
        "得到二值 mask 后按轮廓几何派生 8 个关键点：top_center "
        "（领口）、left/right_shoulder（肩部）、left/right_armpit"
        "（腋下）、left/right_bottom（下摆）、bottom_center。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "8 个关键点的派生不依赖固定比例，而是利用服装轮廓的"
        "几何特征自适应提取：按 y 升序排列轮廓点后，在四个"
        "语义 y 区间内取极值或对称中点。各关键点的搜索区间与"
        "派生方法见表 3-1。",
        first_line_indent=False,
    )
    add_table(
        doc,
        header=["关键点", "搜索区间（相对服装高度 ch）", "派生方法"],
        rows=[
            ["top_center", "[0, 0.25·ch]",
             "顶部 25% 点按 x 中位分左右，各取 y 最小，中点"],
            ["shoulder", "[neck_y+0.12, neck_y+0.22]",
             "区间内取 x 极值"],
            ["armpit", "[shoulder_bottom+0.03, neck_y+0.50]",
             "区间内取 x 极值"],
            ["bottom", "[y_max-0.05·ch, y_max]",
             "区间内取 x 极值 + 内缩 8% 宽度"],
        ],
        col_widths_cm=[3.0, 5.5, 7.5],
    )
    add_caption(doc, "表 3-1  服装 8 关键点的搜索区间与派生方法")

    add_heading(doc, "3.5 流水式 warp 模块", level=2)
    add_paragraph(
        doc,
        "warp_clothing 是入口；method='flow'（默认）走两阶段"
        "路径。整体结构为：Stage A 仿射粗定位 → Stage B 沿 "
        "silhouette 逐行 fit（per-region 独立）→ V 领保留 + "
        "边界羽化。",
        first_line_indent=True,
    )

    add_heading(doc, "3.5.1 Stage A：仿射粗定位", level=3)
    add_paragraph(
        doc,
        "用肩线中点作为对齐基准（服装端 = ClothingDetector 派生"
        "的 left/right_shoulder 中点；人体端 = MediaPipe 派生"
        "的 left/right_shoulder 中点），scale = "
        "min(max(dst_w/src_w, dst_h/src_h) · 1.05, 1.10)："
        "max 保证两维度都不小于身体 bbox，1.05 让衣服略大"
        "防止露肉，1.10 截断防止紧身衣在宽体上算出 1.6x+ 的"
        "过度放大。再用 cv2.warpAffine 把服装 RGB 和 mask 一起"
        "投到人体图坐标系。",
        first_line_indent=True,
    )

    add_heading(doc, "3.5.2 Stage B：按 silhouette 分块 fit", level=3)
    add_paragraph(
        doc,
        "Stage A 只做等比缩放 + 平移，无法适应身体宽度沿 y 的"
        "非线性变化。Stage B 引入“流水式”逐行 fit：首先用 "
        "x/y 极值 + 分位数启发式把身体像素分成 torso / 左袖 / "
        "右袖 三个 mask，袖子 y 范围取腋下到腕（body_pts y 的 "
        "5% 与 85% 分位），外缘逐行夹到 body silhouette。每个"
        "区域独立计算 (s, t) 其中 s=body_w/cloth_w，"
        "t = body_left − s·cloth_left；clamp s 到 [0.7, 1.5] "
        "防止异常值；用 cv2.remap 实现 per-region 拉伸。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "为消除斜襟类衣服在 silhouette 上“左多 X px”的单边"
        "噪声（否则 per-row fit 后会被当成真左边缘拉伸成可见"
        "凸块），把 cloth_left/cloth_right 围绕 body_center "
        "镜像成对称后再参与 (s, t) 计算，dress 宽度与 body-shape "
        "适配不受影响。",
        first_line_indent=True,
    )

    add_heading(doc, "3.5.3 V 领保留与边界羽化", level=3)
    add_paragraph(
        doc,
        "躯干前 3 个采样点 (s, t) 由 (1, 0) 线性插值到过渡点，"
        "避免领口被拉变形。在 torso_x_hi/lo 两侧各 3 像素用 "
        "cosine 渐变对 s、t 做加权，消除 region 拼接处的硬阶跃。",
        first_line_indent=True,
    )

    add_heading(doc, "3.6 融合模块", level=2)
    add_paragraph(
        doc,
        "blend 函数做三步 mask 修正：(1) 用 warped_rgb.max>0 "
        "把 cv2.warpAffine 对 RGB 和 mask 的独立插值导致的多余"
        "外圈 mask 收回；(2) mask 预 erode ~1% min_dim 剔除服装"
        "图自带的“白边”；(3) 高斯模糊 mask 后做 alpha 混合。中间"
        "产物 warped_clothing / warped_mask 用 PNG 无损保存，"
        "避免 JPEG 压缩对 RGB 与 mask 通道处理不一致导致的"
        "“边界偏移”假象。",
        first_line_indent=True,
    )

    add_heading(doc, "3.7 运行结果示例", level=2)
    add_image(doc, SAMPLE_RESULT, width_cm=10)
    add_caption(doc, "图 3-1  系统在单组输入下的试衣合成结果示例")
    add_image(doc, DEBUG_OVERLAY, width_cm=10)
    add_caption(doc, "图 3-2  warp 后服装 mask 叠加到人体的中间结果")
    add_image(doc, CONTACT_SHEET, width_cm=15)
    add_caption(doc, "图 3-3  3 人 × 4 衣 contact sheet 全量组合对比")

    doc.add_page_break()


# ---------- 第四章 实验与结果分析 ----------

def add_chapter_experiment(doc: Document) -> None:
    add_heading(doc, "第四章 实验与结果分析", level=1)

    add_heading(doc, "4.1 测试数据集与运行环境", level=2)
    add_paragraph(
        doc,
        "自建测试集包含 3 张不同身高 / 体型的正面人体图（image / "
        "image2 / image3）与 4 件不同版型的服装图（image：浅粉 T 恤；"
        "image2：蓝旗袍；image3：白衬衫；image4：浅色连衣裙），"
        "共 12 个组合。运行环境为 Python 3.13 + OpenCV 4.8 + "
        "MediaPipe 0.10 + MediaPipe pose_landmarker_full.task 模型，"
        "CPU 推理（未使用 GPU）。",
        first_line_indent=True,
    )

    add_heading(doc, "4.2 试衣结果样例", level=2)
    add_image(doc, SAMPLE_RESULT, width_cm=10)
    add_caption(doc, "图 4-1  单组输入下的试衣合成结果示例")
    add_image(doc, CONTACT_SHEET, width_cm=15)
    add_caption(doc, "图 4-2  12 组组合的 contact sheet 全量对比")

    add_heading(doc, "4.3 各模块对最终结果的影响", level=2)
    add_paragraph(
        doc,
        "本系统的 warp 由 Stage A 与 Stage B 两阶段组成。Stage A "
        "用仿射变换做粗定位，单独运行（method='affine'）即可"
        "得到“整体位置对但形状生硬”的结果：长袖不会贴合到手臂，"
        "身体宽度沿 y 的非线性变化被压平成等比缩放。引入 Stage B "
        "后，躯干/袖子分块独立 fit，结果在腰部收紧、袖口贴合等"
        "位置明显改善。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "在分块基础上，对 cloth silhouette 围绕 body_center 做"
        "对称化处理消除了斜襟衣服带来的“单边凸块”；V 领保留"
        "前 3 个采样点的 (s, t) 由 (1, 0) 插值，避免领口被拉"
        "变形；torso_x_hi/lo 两侧各 3 像素的 cosine 羽化消除了"
        "region 拼接处的垂直接缝。三者联合作用下，长袖 / 旗袍 / "
        "连衣裙的视觉效果均显著优于“仅 Stage A 仿射”。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "融合模块方面，mask 收回、erode、高斯羽化三步缺一不可："
        "不做 mask 收回会出现黑色边框；不做 erode 会透出服装图"
        "的白边；不做高斯羽化会留下硬接缝。三步联合后方可得到"
        "无缝的合成图。",
        first_line_indent=True,
    )

    add_heading(doc, "4.4 失败案例与局限性分析", level=2)
    add_paragraph(
        doc,
        "在 12 组测试中，系统在以下场景表现稳定：正面站姿、"
        "常规上身/连衣裙、浅色或深色背景、人体无严重遮挡。但"
        "在以下场景存在可见缺陷：",
        first_line_indent=True,
    )
    for item in [
        "复杂版型（荷叶边、泡泡袖）：1D 行级 (s, t) 拉伸只能改变"
        "水平方向的覆盖宽度，无法表达 y 方向的体积感。",
        "极端姿态（侧身、坐姿、双手交叉）：body_region_contour "
        "基于 33 关键点的几何多边形在极端姿态下会出现“自交”或"
        "“肩宽估计错误”，导致 Stage A 仿射错位。",
        "严重遮挡（手提包遮胯、长发遮肩）：MediaPipe 关键点"
        "可能被可见性阈值剔除，body_region_contour 退回到"
        "无手臂的旧版多边形，长袖 fit 失败。",
        "过短或过长服装（吊带、过膝长裙）：搜索区间按相对比例 "
        "ch 设置，对极端比例会失配，腋下 / 下摆关键点偏移。",
        "复杂背景服装图（模特自身穿着展示）：CLAHE + Canny "
        "对人物-服装边缘的区分不可靠，mask 可能把人脸一并"
        "抠出。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.5 本章小结", level=2)
    add_paragraph(
        doc,
        "实验表明，本系统在“正面站姿 + 常规版型”这一目标场景"
        "下能稳定产出视觉效果合理的试穿结果，但在复杂版型、"
        "极端姿态、严重遮挡、复杂背景等场景下仍有可见缺陷。"
        "这些缺陷是 1D 行级 warp 与固定几何先验的固有局限，"
        "改进方向将在第六章讨论。",
        first_line_indent=True,
    )

    doc.add_page_break()


# ---------- 第五章 与开源虚拟试衣方法的对比 ----------

def add_chapter_compare(doc: Document) -> None:
    add_heading(doc, "第五章 与开源虚拟试衣方法的对比", level=1)

    add_heading(doc, "5.1 开源方法综述", level=2)

    add_heading(doc, "5.1.1 VITON 与 CP-VTON", level=3)
    add_paragraph(
        doc,
        "VITON [1] 首次用端到端网络做虚拟试衣：先把人体表示为"
        "“与服装无关但保留姿态/外貌”的描述，再用 coarse-to-fine "
        "策略——先生成粗略的试穿图（把目标服装直接叠到人身上），"
        "再用一个 refinement 网络增强细节。VITON 输出 256×192，"
        "纹理细节模糊。CP-VTON [2] 在此基础上引入 GMM（Geometric "
        "Matching Module）+ TPS：用 GMM 把服装薄板样条参数学习"
        "出来，再用 Try-On Module 学到的 composition mask 把 "
        "warp 后的服装与人像融合。",
        first_line_indent=True,
    )

    add_heading(doc, "5.1.2 VITON-HD 与 HR-VITON", level=3)
    add_paragraph(
        doc,
        "VITON-HD [3] 把输出提升到 1024×768：先用分割图引导"
        "服装粗 fit 到身体，再用 ALIAS（ALIgnment-Aware Segment）"
        "归一化 + ALIAS 生成器处理错位区域并保留高分辨率细节。"
        "HR-VITON [4] 进一步把“warp 阶段”和“分割生成阶段”合并"
        "为一个统一的 try-on condition generator，用 feature "
        "fusion 块让两阶段信息互通，避免错位与 pixel-squeezing "
        "伪影，并用 discriminator rejection 过滤错误的分割预测。"
        "两者都需要成对训练数据 + 较强 GPU。",
        first_line_indent=True,
    )

    add_heading(doc, "5.1.3 基于外观流的方法（GFLA / He 等 / GP-VTON）", level=3)
    add_paragraph(
        doc,
        "GFLA [5] 提出“differentiable global-flow local-attention”"
        "框架：先用全局相关性预测 flow 场，再对 flowed local patch "
        "对算 local attention 系数，最后用 content-aware sampling "
        "对源特征做 warp。He 等 [6] 在 CVPR 2022 首次把 StyleGAN "
        "用在 appearance flow 估计上，用全局 style 向量编码全图"
        "上下文，再用 flow refinement 模块补充局部细节。GP-VTON [7] "
        "在 CVPR 2023 提出 Local-Flow Global-Parsing（LFGP）warp "
        "模块 + Dynamic Gradient Truncation（DGT）训练策略，"
        "分部件局部 warp 后用全局 parsing 拼装，避免传统全局 warp "
        "的纹理挤压。本系统的 Stage B “按行 per-region (s, t)”"
        "在思想上是 1D 版 appearance flow——只对 x 方向求解，"
        "速度更快但表达能力受限。",
        first_line_indent=True,
    )

    add_heading(doc, "5.1.4 基于扩散模型的方法（LaDI-VTON / OOTDiffusion / StableVITON）", level=3)
    add_paragraph(
        doc,
        "LaDI-VTON [8] 是首个把 latent diffusion 引入虚拟试衣的"
        "工作：模型是一个带可学习 skip-connection 的 autoencoder "
        "+ 潜在扩散模型；用 textual inversion 把服装的视觉特征"
        "映射到 CLIP token 嵌入空间，得到一组 pseudo-word token "
        "嵌入去条件化生成过程，从而保留纹理细节。OOTDiffusion [9] "
        "（AAAI 2025）用 outfitting UNet 学服装细节，并通过 "
        "outfitting fusion 在去噪 UNet 的 self-attention 层中把"
        "服装特征对齐到目标人体——显式 warp 被完全省略；outfitting "
        "dropout + classifier-free guidance 提供可控强度。StableVITON "
        "[10] 在预训练扩散模型的潜在空间内端到端学习服装—人体的"
        "语义对应，用 zero cross-attention 块同时承担 warp 与细节"
        "保留，配合 attention total variation loss 得到更锐利的"
        "注意力图。该系列对数据量、显存和推理时间的要求显著高于"
        "前几代。",
        first_line_indent=True,
    )

    add_heading(doc, "5.2 多维度对比", level=2)
    add_table(
        doc,
        header=["方法", "发表", "技术路线（按论文 abstract 核实）"],
        rows=[
            ["VITON [1]", "CVPR 2018", "Coarse-to-fine 合成 + refinement network"],
            ["CP-VTON [2]", "ECCV 2018", "GMM 学习的 TPS 变形 + composition mask"],
            ["VITON-HD [3]", "CVPR 2021", "分割引导 + ALIAS 归一化 + ALIAS 生成器"],
            ["HR-VITON [4]", "ECCV 2022", "统一 try-on condition generator + feature fusion"],
            ["GFLA [5]", "arXiv 2020", "Differentiable global-flow + local-attention"],
            ["He et al. [6]", "CVPR 2022", "StyleGAN-based global appearance flow + refinement"],
            ["GP-VTON [7]", "CVPR 2023", "Local-Flow Global-Parsing (LFGP) + DGT 训练"],
            ["LaDI-VTON [8]", "ACM MM 2023", "Latent diffusion + textual inversion (CLIP)"],
            ["OOTDiffusion [9]", "AAAI 2025", "Outfitting fusion in self-attention + outfitting dropout"],
            ["StableVITON [10]", "arXiv 2023", "Zero cross-attention + attention TV loss"],
            ["本系统", "—", "几何 warp + silhouette fit（无训练）"],
        ],
        col_widths_cm=[3.5, 3.0, 9.5],
    )
    add_caption(doc, "表 5-1  本系统与主流开源虚拟试衣方法概览（仅列已核实的论文/会议）")

    add_table(
        doc,
        header=["对比维度", "本系统",
                "早期 TPS 系（VITON/CP-VTON）",
                "高分辨率系（VITON-HD/HR-VITON）",
                "扩散模型系（LaDI/OOT/StableVITON）"],
        rows=[
            ["是否需要训练", "否", "是", "是", "是"],
            ["是否需要成对数据", "否", "是", "是（高分辨率成对数据）",
             "是（成对 + 预训练扩散模型）"],
            ["是否依赖预训练权重", "否（MediaPipe 除外）", "否",
             "否", "是（依赖 Stable Diffusion）"],
            ["推理硬件", "CPU 即可", "GPU", "GPU（建议 ≥ 16 GB 显存）",
             "GPU（建议 ≥ 24 GB 显存）"],
            ["输出分辨率", "原始分辨率", "256×192",
             "1024×768", "512-1024（扩散采样分辨率）"],
            ["形状表达", "1D 行级 (s, t) 拉伸",
             "TPS 薄板样条（参数化网格）",
             "ALIAS 归一化 / 统一 condition generator",
             "全局 self-attention 融合 / zero cross-attention"],
            ["可解释性", "强（每步几何意义）",
             "中（可视化对应点）", "中", "弱（黑盒）"],
            ["跨数据集泛化", "强（无训练）", "弱", "中", "中"],
            ["本地化部署", "强（无外部依赖）", "中", "中", "中"],
            ["适合场景", "教学 / demo / 0 样本 / 隐私敏感",
             "研究基线 / 256×192 试穿", "电商中高质量试穿",
             "电商高端试穿"],
        ],
        col_widths_cm=[3.0, 2.5, 2.8, 2.8, 3.5],
    )
    add_caption(doc, "表 5-2  多维度横向对比（推理速度等具体数值因硬件/实现差异"
                     "较大，未在公开论文中给出统一基准，故未列出）")

    add_heading(doc, "5.3 本系统的定位与适用场景", level=2)
    add_paragraph(
        doc,
        "把上表抽象成一条主线：开源方法都在“用更大模型换更好"
        "视觉”，本系统走的是另一条路——“用几何 + 关键点换零数据 "
        "+ 可解释”。两条路不存在绝对优劣，而是适合不同的工程"
        "约束。",
        first_line_indent=True,
    )
    add_bullet(doc, "教学 / 课程作业场景：本系统可直接展示“如何用"
                    "关键点驱动 warp”，每一步都在 100 行内可解释；"
                    "深度学习方法需要数小时训练 + 复杂环境才能跑通。")
    add_bullet(doc, "0 样本 / 冷启动场景：当没有标注数据、又必须"
                    "快速给一个 demo 时，传统方法是唯一可行方案。")
    add_bullet(doc, "隐私敏感场景：服装与人像都只在本地 CPU 推理，"
                    "不需要把数据上传到任何云端推理服务。")
    add_bullet(doc, "可作为预处理器：本系统的输出（warped_clothing "
                    "+ warped_mask）可直接作为下游 GFLA / DCI-VTON "
                    "的几何先验，进一步减少这些模型需要学习的形变量。")

    add_heading(doc, "5.4 改进方向", level=2)
    for item in [
        "把 1D 行级 (s, t) 升级为 2D appearance flow：在 30 个"
        "采样 y 之外再增加 30 个采样 x，把 silhouette 拓展成 "
        "pixel-wise flow field（对应 GFLA 的简化版）。",
        "引入轻量 CNN 做 silhouette refinement：用 1-2 层 conv "
        "对 body_silhouette_per_row 输出做平滑，解决 y=0.05/0.85 "
        "分位的硬阶跃。",
        "把 Stage A 的 scale 改为 garment-type aware：保留 "
        "shoulder_w API，按“紧身 / 宽松”两档在 1.05-1.10 之间"
        "自适应。",
        "颜色迁移：在 blend 前对 warped_rgb 做 Lab 空间的均值/方差"
        "匹配，减少“塑料感”。",
        "人脸 / 头发遮罩：在 body_pts 之外生成“保留人脸”的 mask，"
        "防止衣服覆盖到不该覆盖的位置。",
        "实时交互：把 Stage A 改成 GPU (cv2.cuda.warpAffine)，"
        "并把 Stage B 的 remap 改成 PyTorch 实现，争取 1080p "
        "30 FPS。",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()


# ---------- 第六章 总结与展望 ----------

def add_chapter_conclusion(doc: Document) -> None:
    add_heading(doc, "第六章 总结与展望", level=1)

    add_heading(doc, "6.1 工作总结", level=2)
    add_paragraph(
        doc,
        "本文实现了一套基于传统计算机视觉方法的虚拟试衣系统，"
        "核心思路是“以关键点驱动几何 warp”。系统采用三级人体"
        "检测降级链 + CLAHE 多通道 Canny 服装分割 + Stage A/B "
        "两阶段 warp + 边界羽化融合，在自建测试集的 12 组正面"
        "站姿组合上稳定产出视觉效果合理的试穿结果。",
        first_line_indent=True,
    )
    add_paragraph(
        doc,
        "与 VITON [1]、CP-VTON [2]、VITON-HD [3]、HR-VITON [4]、"
        "GFLA [5]、He 等 [6]、GP-VTON [7]、LaDI-VTON [8]、"
        "OOTDiffusion [9]、StableVITON [10] 等深度学习方法"
        "相比，本系统在数据需求、硬件依赖、可解释性上有明显"
        "优势，在纹理细节与复杂版型表达上则有差距。这种权衡"
        "决定了本系统更适合“教学 / 0 样本 / 隐私敏感 / 作为"
        "下游模型的几何先验”等场景，而非直接对标电商级别的"
        "高保真试穿。",
        first_line_indent=True,
    )

    add_heading(doc, "6.2 未来工作", level=2)
    add_paragraph(
        doc,
        "未来工作可以沿两条线推进：(1) 把 1D 行级 warp 升级为 "
        "2D appearance flow，进一步提升复杂版型的还原度；"
        "(2) 把本系统的输出作为下游 GFLA / DCI-VTON / "
        "StableVITON 的几何先验，组合传统方法的速度与深度学习"
        "方法的视觉质量。具体的改进方向见 5.4 节。",
        first_line_indent=True,
    )

    add_heading(doc, "参考文献", level=1)
    # GB/T 7714-2015《信息与文献 参考文献著录规则》格式：
    # 期刊 [J]：作者. 题名[J]. 刊名, 年, 卷(期): 页码.
    # 会议 [C]：作者. 题名[C]//会议录名. 出版地: 出版者, 年: 页码.
    # 专著 [M]：作者. 书名[M]. 出版地: 出版社, 年: 页码.
    # 电子文献 [EB/OL]：作者. 题名[EB/OL]. 网址.
    refs = [
        "[1] HAN X, WU Z, WU Z, et al. VITON: an image-based virtual try-on network[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Salt Lake City: IEEE, 2018: 7543-7552.",
        "[2] WANG B, ZHENG H, LIANG X, et al. Toward characteristic-preserving image-based virtual try-on network[C]//European Conference on Computer Vision. Munich: Springer, 2018: 16-32.",
        "[3] CHOI S, PARK S, LEE M, et al. VITON-HD: high-resolution virtual try-on via misalignment-aware normalization[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. Nashville: IEEE, 2021: 14131-14140.",
        "[4] LEE S, GU G, PARK S, et al. High-resolution virtual try-on with misalignment and occlusion-handled conditions[C]//European Conference on Computer Vision. Tel Aviv: Springer, 2022: 204-219.",
        "[5] REN Y, YU X, CHEN J, et al. Deep image spatial transformation for person image generation[EB/OL]. arXiv, 2020. https://arxiv.org/abs/2003.00696.",
        "[6] HE S, SONG Y Z, XIANG T. Style-based global appearance flow for virtual try-on[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. New Orleans: IEEE, 2022: 3470-3479.",
        "[7] XIE Z, HUANG Z, DONG X, et al. GP-VTON: towards general purpose virtual try-on via collaborative local-flow global-parsing learning[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. Vancouver: IEEE, 2023: 5548-5558.",
        "[8] MORELLI D, BALDRATI A, CARTELLA G, et al. LaDI-VTON: latent diffusion textual-inversion enhanced virtual try-on[C]//Proceedings of the ACM International Conference on Multimedia. Ottawa: ACM, 2023: 8580-8589.",
        "[9] XU Y, GU T, CHEN W, et al. OOTDiffusion: outfitting fusion based latent diffusion for controllable virtual try-on[C]//Proceedings of the AAAI Conference on Artificial Intelligence. Philadelphia: AAAI, 2025: 8996-9004.",
        "[10] KIM J, GU G, PARK M, et al. StableVITON: learning semantic correspondence with latent diffusion model for virtual try-on[EB/OL]. arXiv, 2023. https://arxiv.org/abs/2312.01725.",
        "[11] LUGARESI C, TANG J, NASH H, et al. MediaPipe pose: real-time multi-person pose estimation[EB/OL]. https://google.github.io/mediapipe/solutions/pose.html. (2024-09-01)[2026-06-20].",
        "[12] BRADSKI G. The OpenCV library[J]. Dr. Dobb's Journal of Software Tools, 2000, 25(11): 120-125.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        run = p.add_run(r)
        set_cn_font(run, size=10)


# ---------- 主流程 ----------

def _set_run_font_on_style(style, font_name: str, size_pt: int = 11,
                           bold: bool = False, color_hex: str = "000000") -> None:
    """在 style 的 rPr 上写字体/字号/颜色，绕过 BaseStyle.font 的类型盲区。"""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = rPr.makeelement(qn("w:sz"), {})
        rPr.append(sz)
    sz.set(qn("w:val"), str(size_pt * 2))
    color = rPr.find(qn("w:color"))
    if color is None:
        color = rPr.makeelement(qn("w:color"), {})
        rPr.append(color)
    color.set(qn("w:val"), color_hex)
    if bold:
        b = rPr.find(qn("w:b"))
        if b is None:
            b = rPr.makeelement(qn("w:b"), {})
            rPr.append(b)


def build(out_path: Path) -> None:
    doc = Document()

    _set_run_font_on_style(doc.styles["Normal"], CN_FONT,
                           size_pt=11, color_hex="000000")
    _set_run_font_on_style(doc.styles["Heading 1"], CN_FONT,
                           size_pt=18, bold=True, color_hex="000000")
    _set_run_font_on_style(doc.styles["Heading 2"], CN_FONT,
                           size_pt=15, bold=True, color_hex="000000")
    _set_run_font_on_style(doc.styles["Heading 3"], CN_FONT,
                           size_pt=12, bold=True, color_hex="000000")

    add_cover(doc)
    add_abstract(doc)
    add_toc_placeholder(doc)
    add_chapter_intro(doc)
    add_chapter_background(doc)
    add_chapter_design(doc)
    add_chapter_experiment(doc)
    add_chapter_compare(doc)
    add_chapter_conclusion(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] 报告已写入 {out_path}  ({out_path.stat().st_size/1024:.1f} KB)")


def main() -> int:
    ap = argparse.ArgumentParser(description="生成虚拟试衣系统 Word 报告")
    ap.add_argument("--output", default=str(DEFAULT_OUT),
                    help="输出 .docx 路径（默认 docs/report.docx）")
    args = ap.parse_args()
    build(Path(args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())